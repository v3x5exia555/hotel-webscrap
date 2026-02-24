import sys
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx(source_md_path, output_docx_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('Pahang Hotel Intelligence: Quick Reference', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Read the MD file
    with open(source_md_path, 'r') as f:
        md_content = f.read()

    # Simple MD to Docx logic
    lines = md_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|'):
            # Basic table detection - skip raw pipe lines as we handle the table manually below
            continue
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
                    
    # Add the Risk Matrix manually for better formatting
    doc.add_heading('Risk Matrix & Action Plan', level=1)
    
    table_data = [
        ['Priority', 'Scenario', 'Action'],
        ['CRITICAL', 'Unregistered + Active Sales', 'Enforcement: Mandatory Registration'],
        ['HIGH', 'Registered + Large Sales Gap', 'Audit: Review financial records'],
        ['WARN', 'Minor reporting gaps', 'Monitor: Send automated reminder'],
        ['OK', 'Data matches reports', 'Compliance: No action needed']
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = table_data[0][0]
    hdr_cells[1].text = table_data[0][1]
    hdr_cells[2].text = table_data[0][2]
    
    for p, s, a in table_data[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = p
        row_cells[1].text = s
        row_cells[2].text = a

    doc.save(output_docx_path)
    print(f"File saved successfully to {output_docx_path}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description='Convert Markdown to Docx')
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('output', help='Output Docx file')
    args = parser.parse_args()
    
    create_docx(args.input, args.output)
